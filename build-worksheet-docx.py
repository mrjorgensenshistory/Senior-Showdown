"""
Builds worksheet.docx and worksheet-answer-key.docx from the same content
as the HTML versions. Run with: python build-worksheet-docx.py
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, color_hex):
    """Set background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top, bottom, left, right -> color hex string."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '8')
            el.set(qn('w:color'), kwargs[edge])
            borders.append(el)


def add_section_header(doc, text, bg='000000', fg='FFFFFF'):
    """Black bar with white text — for PART 1, PART 2, etc."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, bg)
    cell.width = Inches(7.0)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(fg)
    return table


def add_grade_band(doc, text):
    """Green band with gold text — for 9TH GRADE, 10TH GRADE, etc."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_background(cell, '1E5A2E')
    cell.width = Inches(7.0)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string('D4AF37')
    return table


def add_question(doc, number, text, blank_after='_______', blank_size='medium'):
    """Add a numbered question with fill-in-the-blank."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(f"{number}. {text} ")
    run.font.size = Pt(10)
    if blank_after:
        run2 = p.add_run(blank_after)
        run2.font.size = Pt(10)


def add_writing_lines(doc, count=2, indent=0.4):
    """Add solid horizontal lines for reflection answers."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run('_' * 80)
        run.font.size = Pt(10)


# ============================================================================
# STUDENT WORKSHEET
# ============================================================================

def build_worksheet():
    doc = Document()

    # Letter page with narrow margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # === HEADER ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run('SENIOR SHOWDOWN — PROOF OF COMPLETION')
    run.font.size = Pt(18)
    run.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(6)
    run = sub.add_run('ASB Budget Simulation — You must play all 4 grades to finish this worksheet')
    run.font.size = Pt(10)
    run.font.italic = True

    # Horizontal rule
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'double')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # === DIRECTIONS BOX ===
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(7.5)
    set_cell_background(cell, 'FFF8DC')
    set_cell_border(cell, top='8B6914', bottom='8B6914', left='8B6914', right='8B6914')
    p = cell.paragraphs[0]
    run = p.add_run('Directions: ')
    run.font.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run('Play through all 4 grades (9th, 10th, 11th, 12th). After each grade, the game will give you a code — write the full code in the matching box below. Then answer every question.')
    run2.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # === NAME / PERIOD / DATE ROW ===
    t = doc.add_table(rows=1, cols=3)
    t.autofit = False
    widths = [Inches(4.0), Inches(1.5), Inches(2.0)]
    labels = ['NAME', 'PERIOD', 'DATE']
    for i, label in enumerate(labels):
        c = t.cell(0, i)
        c.width = widths[i]
        set_cell_border(c, bottom='000000')
        p = c.paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(8)
        run.font.bold = True
        # Add a blank line to give writing room above the border
        p2 = c.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.add_run(' ')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # === CODE BOXES (4 grades) ===
    code_table = doc.add_table(rows=1, cols=4)
    code_table.autofit = False
    grade_info = [
        ('9TH GRADE CODE', 'starts with WAVE-'),
        ('10TH GRADE CODE', 'starts with HAUN-'),
        ('11TH GRADE CODE', 'starts with HOCO-'),
        ('12TH GRADE CODE', 'starts with PROM-'),
    ]
    for i, (label, sublabel) in enumerate(grade_info):
        c = code_table.cell(0, i)
        c.width = Inches(1.875)
        set_cell_background(c, 'F4ECD8')
        set_cell_border(c, top='1E5A2E', bottom='1E5A2E', left='1E5A2E', right='1E5A2E')
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        p1 = c.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_after = Pt(0)
        r = p1.add_run(label)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string('1E5A2E')

        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        r = p2.add_run(sublabel)
        r.font.size = Pt(8)
        r.font.italic = True

        # Code blanks: ____-____-___
        p3 = c.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(0)
        r = p3.add_run('____  -  ____  -  ___')
        r.font.size = Pt(10)
        r.font.name = 'Courier New'
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # === PART 1 — GAME KNOWLEDGE ===
    add_section_header(doc, '  PART 1 — GAME KNOWLEDGE (only knowable by playing)')

    add_grade_band(doc, '  9TH GRADE')
    add_question(doc, 1, 'What football character TACKLES you when you skip the paperwork?', '_____________________')
    add_question(doc, 2, 'How many students must COUNT the cash from a fundraiser (per the game)?', '______')
    add_question(doc, 3, 'What is the PROFIT CAP on the 9th-grade dance in this game?', '$ ______')
    add_question(doc, 4, 'Adding food to the 9th-grade dance boosts attendance by what %?', '______ %')
    add_question(doc, 5, 'About how many students attend ACE total (in the game)?', '______')

    add_grade_band(doc, '  10TH GRADE')
    add_question(doc, 6, 'The biggest October event 10th-graders run is the', '_____________________')
    add_question(doc, 7, 'If you charge MORE than $30 for a dance ticket and DON’T serve food, what happens?', '_____________________')
    add_question(doc, 8, 'Max cars you can wash in a 3-hour car wash:', '______ cars')
    add_question(doc, 9, 'How much do car wash SUPPLIES cost?', '$ ______')
    add_question(doc, 10, 'What month is the Bundt Cake fundraiser?', '_____________________')

    add_grade_band(doc, '  11TH GRADE')
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('11. Name TWO cutscene characters in the game (not students):')
    r.font.size = Pt(10)
    for letter in ['a', 'b']:
        sub = doc.add_paragraph()
        sub.paragraph_format.left_indent = Inches(0.5)
        sub.paragraph_format.space_after = Pt(2)
        r = sub.add_run(f'{letter}) _____________________________')
        r.font.size = Pt(10)
    add_question(doc, 12, 'What PERCENTAGE of HOCO profit does the JUNIOR class get in this game?', '______ %')
    add_question(doc, 13, 'A typical HOCO net profit in this game is roughly:', '$ ______')
    add_question(doc, 14, 'The "Back to the Future"-themed prom venue is named:', '_____________________________')

    add_grade_band(doc, '  12TH GRADE')
    add_question(doc, 15, 'How much does the Disney Grad Nite bus cost?', '$ ______')
    add_question(doc, 16, 'Yearbook cost per student:', '$ ______')
    add_question(doc, 17, 'The FREE prom venue is named:', '_____________________________')
    add_question(doc, 18, 'The MOST EXPENSIVE prom venue is named:', '_____________________________')
    add_question(doc, 19, 'Football character who gets DISAPPOINTED when fundraisers fail:', '_____________________')
    add_question(doc, 20, 'How much does the PRO DJ cost (the Student DJ is free)?', '$ ______')

    # === PAGE 2 (force break onto next paragraph; no extra empty paragraph) ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    title.paragraph_format.page_break_before = True
    run = title.add_run('SENIOR SHOWDOWN — PAGE 2')
    run.font.size = Pt(14)
    run.font.bold = True

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'double')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Name / Period row again
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    widths = [Inches(5.5), Inches(2.0)]
    labels = ['NAME (AGAIN)', 'PERIOD']
    for i, label in enumerate(labels):
        c = t.cell(0, i)
        c.width = widths[i]
        set_cell_border(c, bottom='000000')
        p = c.paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(8)
        run.font.bold = True
        c.add_paragraph(' ').paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # === PART 2 — YOUR PLAYTHROUGH ===
    add_section_header(doc, '  PART 2 — YOUR SPECIFIC PLAYTHROUGH')
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run('  These answers will be DIFFERENT for every student. Copy directly from your game.')
    r.font.italic = True
    r.font.size = Pt(9)

    add_question(doc, 21, '9th grade ENDING balance:', '$ ______________')
    add_question(doc, 22, '10th grade ENDING balance:', '$ ______________')
    add_question(doc, 23, '11th grade ENDING balance:', '$ ______________')
    add_question(doc, 24, '12th grade ENDING balance:', '$ ______________')
    add_question(doc, 25, 'Which prom venue did YOU pick for senior year?', '_____________________________')
    add_question(doc, 26, 'Did YOU pick Pro DJ or Student DJ for your dance?', '_____________________')
    add_question(doc, 27, 'Did YOU include food at the 9th-grade dance? (YES / NO)', '______')
    add_question(doc, 28, 'What ticket price did YOU set for your 9th-grade dance?', '$ ______')

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('29. List THREE fundraisers YOU ran in 12th grade (in order):')
    r.font.size = Pt(10)
    for n in [1, 2, 3]:
        sub = doc.add_paragraph()
        sub.paragraph_format.left_indent = Inches(0.5)
        sub.paragraph_format.space_after = Pt(2)
        r = sub.add_run(f'{n}) _________________________________________________')
        r.font.size = Pt(10)

    add_question(doc, 30, 'What price did YOU set for the senior PACKAGE (prom + Grad Nite + yearbook)?', '$ ______________')

    # === PART 3 — REFLECTION ===
    add_section_header(doc, '  PART 3 — REFLECTION (your game + real ASB)')
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(4)
    r = sub.add_run('  1–2 sentences each. Reference specifics from YOUR playthrough AND from real life at our school.')
    r.font.italic = True
    r.font.size = Pt(9)

    reflections = [
        ('31', 'Which fundraiser YOU ran made the MOST money in the game? About how much profit?'),
        ('32', 'Describe a fundraiser you regret — one that lost money, got cancelled, or you skipped. What happened?'),
        ('33', 'Compare this game to REAL ASB at our school. What did it get RIGHT? Name something specific.'),
        ('34', 'What did the game get WRONG or LEAVE OUT about real ASB at our school? Be specific.'),
        ('35', 'What is the MOST important thing you LEARNED from this game that applies to real ASB or class money?'),
    ]
    for num, q in reflections:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f'{num}. {q}')
        r.font.size = Pt(10)
        add_writing_lines(doc, count=2, indent=0.3)

    # === TEACHER FOOTER ===
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    t = doc.add_table(rows=2, cols=3)
    t.autofit = False
    for c in t.row_cells(0) + t.row_cells(1):
        set_cell_border(c, top='000000', bottom='000000', left='000000', right='000000')

    t.cell(0, 0).merge(t.cell(0, 2))
    p = t.cell(0, 0).paragraphs[0]
    r = p.add_run('TEACHER USE ONLY')
    r.font.bold = True
    r.font.size = Pt(9)

    labels = ['All 4 codes match master list?  YES / NO', 'Avg quiz score', 'Grade / Notes']
    for i, label in enumerate(labels):
        c = t.cell(1, i)
        p = c.paragraphs[0]
        r = p.add_run(label)
        r.font.size = Pt(8)
        r.font.bold = True
        c.add_paragraph('_______________________').paragraph_format.space_after = Pt(0)

    doc.save('worksheet.docx')
    print('Wrote worksheet.docx')


if __name__ == '__main__':
    build_worksheet()

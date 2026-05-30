"""
Build standalone DOCX + PDF for the 3 ASB forms baked into the game:
  1. Fundraising Event Request
  2. Expenditure Approval (PO / pre-approval)
  3. Check Request (reimbursement)

Same approach as build-worksheet-docx.py. Run with: python build-forms-docx.py
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), kwargs.get('size', '8'))
            el.set(qn('w:color'), kwargs[edge])
            borders.append(el)


def setup_page(doc):
    """Letter, 0.5" margins."""
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)


def add_header(doc, title, school1, school2, subtitle):
    """Form header — title, school, subtitle, separator."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string('1A3A6B')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(school1)
    r.font.size = Pt(10)

    if school2:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(school2)
        r.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(subtitle)
    r.font.size = Pt(9)
    r.font.italic = True

    # Horizontal rule
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), '1A3A6B')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_rule_box(doc, text, color='8B0000'):
    """Bold warning box."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    set_cell_background(cell, 'FFF8F0')
    set_cell_border(cell, left=color, size='18')
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string('1A3A6B')


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string('555555')


def add_field_row(doc, fields):
    """Single row of inline fields. fields = list of (label, blank_width_inches)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    for i, (label, width) in enumerate(fields):
        if i > 0:
            p.add_run('   ').font.size = Pt(10)
        r = p.add_run(label + ' ')
        r.font.size = Pt(10)
        r.font.bold = True
        # Underscores for the blank
        blank_chars = int(width * 12)  # ~12 underscores per inch
        r = p.add_run('_' * blank_chars)
        r.font.size = Pt(10)


def add_multiline_field(doc, label, lines=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    r.font.size = Pt(10)
    r.font.bold = True
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run('_' * 95)
        r.font.size = Pt(10)


def add_signature_line(doc, label, name_width=4.5, with_date=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + ' ')
    r.font.size = Pt(10)
    r.font.bold = True
    name_blanks = int(name_width * 12)
    r = p.add_run('_' * name_blanks)
    r.font.size = Pt(10)
    if with_date:
        r = p.add_run('   Date: ')
        r.font.size = Pt(10)
        r.font.bold = True
        r = p.add_run('_' * 12)
        r.font.size = Pt(10)


# ============================================================================
# FORM 1 — FUNDRAISING EVENT REQUEST
# ============================================================================
def build_fundraising_form():
    doc = Document()
    setup_page(doc)

    add_header(
        doc,
        'REQUEST FOR APPROVAL: FUNDRAISING EVENT',
        'Academy of Careers and Exploration',
        'Helendale Secondary School — Helendale, CA',
        'Applications must be submitted at least four (4) weeks before the fundraiser.',
    )

    add_rule_box(
        doc,
        'No form, no fundraiser. Every form needs three signatures: student, advisor, and principal. The principal has FINAL say.',
    )

    add_field_row(doc, [('Name of School:', 3.5), ('Date Submitted:', 1.5)])
    add_field_row(doc, [('Name of Club / Class:', 5.0)])
    add_field_row(doc, [('Name of Person Completing Form:', 4.0)])
    add_field_row(doc, [('Proposed Event:', 5.5)])
    add_field_row(doc, [('Proposed Date(s) of Event:', 3.5)])
    add_field_row(doc, [('Time of Sale:', 4.5)])
    add_multiline_field(doc, 'Description of Fundraiser:', lines=2)
    add_field_row(doc, [('Location:', 5.0)])
    add_field_row(doc, [('Club Contact Person:', 3.0), ('ASB / Club Advisor:', 2.5)])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Status:   ')
    r.font.size = Pt(10)
    r.font.bold = True
    r = p.add_run('☐ New Event       ☐ Held Previously')
    r.font.size = Pt(10)

    # FCMAT food rule callout
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    set_cell_background(cell, 'FFF8DC')
    set_cell_border(cell, top='8B6914', bottom='8B6914', left='8B6914', right='8B6914')
    p = cell.paragraphs[0]
    r = p.add_run('FCMAT FOOD RULE: ')
    r.font.size = Pt(9)
    r.font.bold = True
    r = p.add_run('Food may only be sold during lunch and after school. Not before school. Not during class. Money MAY be collected during school hours.')
    r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    add_section_title(doc, 'CERTIFICATION & APPROVAL SIGNATURES')
    add_note(doc, 'I certify that the information provided is accurate and that the fundraiser will comply with all ASB and school policies.')
    add_signature_line(doc, 'Student Club Representative:')
    add_signature_line(doc, 'Club Advisor:                       ')
    add_signature_line(doc, 'Principal:                          ')
    add_signature_line(doc, 'ASB Bookkeeper:               ')

    doc.save('form-fundraising-request.docx')
    print('Wrote form-fundraising-request.docx')


# ============================================================================
# FORM 2 — EXPENDITURE APPROVAL (PO / pre-approval)
# ============================================================================
def build_expenditure_form():
    doc = Document()
    setup_page(doc)

    add_header(
        doc,
        'EXPENDITURE APPROVAL',
        'Academy of Careers and Exploration / Riverview Middle School',
        '',
        'For items to be purchased without a purchase order, but requiring pre-approval from ASB.',
    )

    add_rule_box(
        doc,
        'PRE-APPROVAL ONLY. The advisor cannot front money for any purchase without this form signed FIRST. No form = no purchase = no reimbursement.',
    )

    add_field_row(doc, [('Date Requested:', 2.5), ('Date of Event:', 2.5)])
    add_field_row(doc, [('Club to be Charged:', 5.0)])
    add_field_row(doc, [('Payee — Name of Club / Staff Member to be Reimbursed:', 3.5)])

    add_section_title(doc, 'LINE ITEMS')

    # Line items table — 5 rows
    items_table = doc.add_table(rows=6, cols=5)
    items_table.autofit = False
    col_widths = [Inches(0.5), Inches(3.5), Inches(0.9), Inches(1.0), Inches(1.2)]
    headers = ['Item #', 'Description', 'Quantity', 'Unit Price', 'Total Amount']

    for i, h in enumerate(headers):
        cell = items_table.cell(0, i)
        cell.width = col_widths[i]
        set_cell_background(cell, 'F0F0F0')
        set_cell_border(cell, top='000000', bottom='000000', left='000000', right='000000')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(9)

    for row in range(1, 6):
        for col in range(5):
            cell = items_table.cell(row, col)
            cell.width = col_widths[col]
            set_cell_border(cell, top='000000', bottom='000000', left='000000', right='000000')
            if col == 0:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(row))
                r.font.bold = True
                r.font.size = Pt(9)
            else:
                # Add a non-breaking space so the cell has at least one character of height
                p = cell.paragraphs[0]
                p.add_run(' ')

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Totals (right-aligned simple rows)
    for label in ['Subtotal:  $', '+ Est. Sales Tax:  $', '+ Est. Shipping:  $']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label + ' ')
        r.font.size = Pt(10)
        r.font.bold = True
        r = p.add_run('_' * 14)
        r.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('TOTAL:  $ ')
    r.font.size = Pt(12)
    r.font.bold = True
    r = p.add_run('_' * 14)
    r.font.size = Pt(12)
    r.font.bold = True

    add_section_title(doc, 'CERTIFICATION & APPROVAL SIGNATURES')
    add_note(doc, 'I certify that these expenditures are necessary for the approved fundraiser and comply with ASB policies.')
    add_signature_line(doc, 'Student Club Representative:')
    add_signature_line(doc, 'Club Advisor:                       ')
    add_signature_line(doc, 'Principal:                          ')
    add_signature_line(doc, 'ASB Bookkeeper:               ')

    doc.save('form-expenditure-approval.docx')
    print('Wrote form-expenditure-approval.docx')


# ============================================================================
# FORM 3 — CHECK REQUEST (close-out / reimbursement)
# ============================================================================
def build_check_request_form():
    doc = Document()
    setup_page(doc)

    add_header(
        doc,
        'ASB CHECK REQUEST FORM',
        'Academy of Careers and Exploration',
        '',
        'Close-out & Reimbursement',
    )

    add_rule_box(
        doc,
        'NO RECEIPT = NO REIMBURSEMENT. Always attach the receipt to this form. The check is issued to the person who fronted the money, not the store.',
    )

    add_field_row(doc, [('Fiscal Year:', 2.5), ('Date Submitted:', 2.5)])
    add_field_row(doc, [('Date of Expenditure:', 2.5), ('Amount (from receipt):  $', 1.5)])
    add_field_row(doc, [('Payee (who gets the check):', 4.5)])
    add_field_row(doc, [('Club / Group:', 5.0)])
    add_multiline_field(doc, 'Purpose of Expenditure:', lines=2)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Receipt Attached?   ')
    r.font.size = Pt(10)
    r.font.bold = True
    r = p.add_run('☐ YES — receipt is attached       ☐ NO — receipt lost')
    r.font.size = Pt(10)

    add_section_title(doc, 'RECONCILIATION (Fundraiser Close-Out)')
    add_field_row(doc, [('Total Revenue Collected:  $', 2.5), ('Total Expenses (from receipt):  $', 2.0)])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Net Profit (Revenue − Expenses):  $ ')
    r.font.size = Pt(11)
    r.font.bold = True
    r = p.add_run('_' * 15)
    r.font.size = Pt(11)
    r.font.bold = True

    add_section_title(doc, 'APPROVAL SIGNATURES')
    add_signature_line(doc, 'ASB / Club Representative:')
    add_signature_line(doc, 'Club Advisor:                       ')
    add_signature_line(doc, 'Principal:                          ')
    add_signature_line(doc, 'ASB Bookkeeper:               ')

    doc.save('form-check-request.docx')
    print('Wrote form-check-request.docx')


if __name__ == '__main__':
    build_fundraising_form()
    build_expenditure_form()
    build_check_request_form()
